from flask import Flask, render_template, request, redirect, url_for, send_file,session,flash, jsonify
from google.cloud import vision
import os
from docx import Document
import io
from werkzeug.utils import secure_filename
from flask_sqlalchemy import SQLAlchemy
from datetime import timedelta, datetime
import bcrypt
from pdf2image import convert_from_path
import pytesseract
import msoffcrypto
import docx
import tempfile
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import Mail
from dotenv import load_dotenv
import re

load_dotenv('sendgrid.env')

app = Flask(__name__,static_folder='static')
app.secret_key = "abcd"
app.config['SQLALCHEMY_DATABASE_URI'] = 'postgresql://postgres:Gang$hit101@localhost/OCR_PROJECT_ORGANIZATION'
app.permanent_session_lifetime = timedelta(days=365)
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024 #16MB
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif','webp'}

db = SQLAlchemy(app)

# Create upload folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)


#CHECK IF FILE TYPE IS ALLOWED
def allowed_file(filename):
    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

#PICS DB CLASS
class org_pictures(db.Model):
    __tablename__ = 'org_pictures'
    created_at = db.Column(db.DateTime, default=datetime.now())
    id = db.Column(db.Integer, primary_key=True)
    image_path = db.Column(db.String(255))
    user_id = db.Column(db.Integer, db.ForeignKey('org_users.id', ondelete='CASCADE'))
    is_public = db.Column(db.Boolean, default=False)


    def __init__(self, image_path=None,is_public = False):
        self.image_path = image_path
        self.is_public = is_public


#USERS DB CLASS
class org_users(db.Model):
    __tablename__ = 'org_users'
    id = db.Column(db.Integer, primary_key=True)
    org_username = db.Column(db.String(400))
    org_password_hash = db.Column(db.String(400))
    email = db.Column(db.String(400))
    created_at = db.Column(db.DateTime, default=datetime.now())
    image_path = db.Column(db.String(400))
    uploader = db.relationship('org_pictures', backref='shared_by_user')
    images = db.relationship('org_pictures', backref='user', lazy=True)
    
    def __init__(self, org_username, org_password_hash, email,image_path):
        self.org_username = org_username
        self.org_password_hash = org_password_hash
        self.email = email
        self.image_path = image_path
    
#PICTURE SHARE CLASS
class PictureShare(db.Model):
    __tablename__ = 'picture_shares'
    id = db.Column(db.Integer, primary_key=True)
    picture_id = db.Column(db.Integer, db.ForeignKey('org_pictures.id', ondelete='CASCADE'))
    shared_with_user_id = db.Column(db.Integer, db.ForeignKey('org_users.id', ondelete='CASCADE'))

    picture = db.relationship('org_pictures', backref='shared_with')
    shared_with_user = db.relationship('org_users', backref='shared_pictures')

    def __init__(self, picture_id, shared_with_user_id):
        self.picture_id = picture_id
        self.shared_with_user_id = shared_with_user_id


#DEFAULT ROUTE
@app.route("/")
def index():
    return render_template("index.html")


#REGISTRATION ROUTE WITH PASSWORD HASHING
@app.route("/submit", methods=["POST", "GET"])
def submit():
    if request.method == "POST":
        org_username = request.form["reg-username"].strip()
        if not re.match(r'^[A-Za-z0-9_]{3,20}$', org_username):
            flash("Username must be 3–20 characters, letters/numbers/underscores only.")
            return redirect(url_for("register"))

        org_password_hash = request.form["reg-password"]
        if len(org_password_hash) < 6:
            flash("Password must be at least 6 characters.")
            return redirect(url_for("register"))

        email = request.form["email"].strip()
        if not re.match(r'^[^@]+@[^@]+\.[^@]+$', email):
            flash("Please enter a valid email address.")
            return redirect(url_for("register"))


        org_password_hash_encode = org_password_hash.encode('utf-8')
        org_password_hash_hashed= bcrypt.hashpw(org_password_hash_encode,bcrypt.gensalt()).decode('utf-8')


        image_path = None


        student1 = org_users(org_username=org_username, org_password_hash=org_password_hash_hashed, email=email, image_path=image_path)

        db.session.add(student1)
        db.session.commit()
        flash("Registration success")

        # 🔥 Send "Thank you for registering" email
        send_thank_you_email(email,org_username)

        return render_template("login.html")

# SEND EMAIL WITH SENDRID
def send_thank_you_email(to_email,username):
    message = Mail(
        from_email='dococrateservices@gmail.com',
        to_emails=to_email,
        subject='SUCESSFUL REGISTRATION',
        html_content=f'<strong>{username} WELCOME TO APP</strong>')
    try:
        sg = SendGridAPIClient('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')
        response = sg.send(message)
        print(response.status_code)
        print(response.body)
        print(response.headers)
    except Exception as e:
        print(e)

#LOGIN ROUTE AND UPLOAD.HTML REDIRECT
@app.route("/search", methods=["POST", "GET"])
def search():
    if request.method == "POST":
        log_user = request.form["requsername"]
        log_password = request.form["reqemail"].encode('utf-8')

        stored_db_password = org_users.query.filter_by(org_username=log_user).first()

        if stored_db_password is None or stored_db_password.org_password_hash is None:
            return render_template("login.html", failed="no user found")


        else:
            stored_db_password2 = stored_db_password.org_password_hash.encode('utf-8')
            isvalid = bcrypt.checkpw(log_password,stored_db_password2)

            found_all_users = org_users.query.filter_by(org_username=log_user).all()
            if isvalid and found_all_users:
                    session["username"] = log_user
                    return render_template("ocr_upload.html",users=found_all_users)
            else:
                    return render_template("login.html", failed="no user found")

    else:
        if "username" in session:
            flash("Already logged in")
            return render_template("ocr_upload.html")
        return render_template("login.html")

#UPLOAD ROUTE TO SAVE IMAGES TO DATABASE WITH USER REDIRECT
@app.route("/upload", methods=["POST", "GET"])
def upload():
    if "username" in session:
        if request.method == "POST":
            image = request.files.get("image")
            image_path = None

            if image and image.filename != "" and allowed_file(image.filename):
                username = session.get("username") or "user"
                #CHANGED SECURE FILENAME
                #filename = secure_filename(f"{username}_{image.filename}")
                filename = secure_filename(f"{image.filename}")
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                image.save(filepath)
                image_path = filename

                # Get the user from DB
                session_user = org_users.query.filter_by(org_username=username).first()
                if session_user:
                    # Always add image to org_pictures
                    new_image = org_pictures(image_path=image_path)
                    session_user.images.append(new_image)

                    # Only set image_path if this is the first upload
                    if not session_user.image_path:
                        session_user.image_path = image_path

                    db.session.add(session_user)
                    db.session.commit()

                    found_all_users = org_users.query.filter_by(org_username=username).first()
                    print_user_pics = found_all_users.id
                    found_all_users_id=org_pictures.query.filter_by(user_id=print_user_pics).all()
                    return render_template("user.html", users=found_all_users_id,username=username)
                else:
                    flash("User not found.")
                    return redirect("/")

            else:
                flash("Invalid or missing image file.")
                return render_template("ocr_upload.html")
        else:
            return render_template("ocr_upload.html")

    else:
        flash("Please log in.")
        return redirect("/login")

#LOGOUT ROUTE TO BE MODIFIED LATER ON
@app.route("/logout")
def logout():
    if "username" in session:
        user = session["username"]
        flash(f"YOU HAVE BEEN LOGGED OUT {user}")
    session.pop("username", None)
    return redirect(url_for("search"))

#ORGANIZATIONAL SYSTEM DOWLOAD WITH LOCK CAPABILITY
@app.route('/download', methods=['POST'])
def download():
    if request.method == "POST":
            edited_text = request.form['real-text']
            protect = request.form.get('lock-doc')
            doc_name = request.form.get('download-doc-name')

            if doc_name:
                doc_name = doc_name
            else:
                doc_name = 'ocr_result'

            if protect:
                    doc_lock = request.form['lock-password']
                    password =  doc_lock # you can also let the user pick the password!
                    protected_docx = create_protected_docx(edited_text, password)

                    return send_file(
                        protected_docx,
                        as_attachment=True,
                        download_name=f"{doc_name}.locked.docx",
                        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                # Create a new document with the text
                doc = Document()
                doc.add_paragraph(edited_text)

                # Save to memory
                file_stream = io.BytesIO()
                doc.save(file_stream)
                file_stream.seek(0)

                # Send the file for download
                return send_file(
                    file_stream,
                    as_attachment=True,
                    download_name=f'{doc_name}.docx',
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

#LOCK PDF FUNCTION
def create_protected_docx(text, password):
        # Step 1: Create a normal docx in memory
        file_stream = io.BytesIO()
        doc = docx.Document()
        doc.add_paragraph(text)
        doc.save(file_stream)

        # Step 2: Protect it with a password
        file_stream.seek(0)
        office_file = msoffcrypto.OfficeFile(file_stream)

        protected_stream = io.BytesIO()
        office_file.encrypt(password=password, outfile=protected_stream)

        protected_stream.seek(0)
        return protected_stream

#UPLOAD GOOGLE VISION OCR FILE INPUT
@app.route('/uploadocr', methods=['POST'])
def upload_image():
    if 'image' not in request.files:
        return 'No image uploaded', 400

    image_file = request.files['image']
    content = image_file.read()

    client = vision.ImageAnnotatorClient()
    image = vision.Image(content=content)

    # Use document_text_detection for structured OCR
    response = client.document_text_detection(image=image)
    annotation = response.full_text_annotation

    # Logo detection (optional)
    logo_response = client.logo_detection(image=image)
    logos = logo_response.logo_annotations
    detected_logos = [logo.description for logo in logos]

    doc = Document()
    doc.add_paragraph("Detected Text:")

    for page in annotation.pages:
        for block in page.blocks:
            block_text = ""
            for paragraph in block.paragraphs:
                para_text = ""
                for word in paragraph.words:
                    word_text = ''.join([
                        symbol.text for symbol in word.symbols
                    ])
                    para_text += word_text + " "
                block_text += para_text.strip() + "\n"

            # This heuristic assumes short lines aligned closely may be a table row
            if len(block.paragraphs) > 1 and all(len(p.words) > 1 for p in block.paragraphs):
                doc.add_paragraph("\n[Possible Table Detected Below]", style='Intense Quote')
                table = doc.add_table(rows=0, cols=1)
                for paragraph in block.paragraphs:
                    row_cells = [ ''.join(s.text for w in paragraph.words for s in w.symbols) ]
                    row = table.add_row().cells
                    row[0].text = row_cells[0]
            else:
                doc.add_paragraph(block_text.strip())

    if detected_logos:
        doc.add_paragraph("\nDetected Logos:")
        for logo in detected_logos:
            doc.add_paragraph(f"- {logo}")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return render_template('result.html',
                           detected_text=annotation.text,
                           detected_logos=detected_logos,
                           file_stream=file_stream)

#UPLOAD GOOGLE VISION USING PERFORM OCR BUTTON
@app.route('/uploadocr2', methods=['POST'])
def upload_image2():
    image_path = request.form.get("image_path")

    if image_path:

        full_path = os.path.join("static", "uploads", image_path)

        try:
            with io.open(full_path, 'rb') as image_file:
                content = image_file.read()
        except FileNotFoundError:
            return "Image file not found."


        client = vision.ImageAnnotatorClient()
        image = vision.Image(content=content)

        response = client.text_detection(image=image)
        texts = response.text_annotations

        if texts:
            detected_text = texts[0].description

            # Create a temporary DOCX file in memory
            doc = Document()
            doc.add_paragraph(detected_text)

            # Save to memory instead of disk
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            return render_template('result.html',
                                   detected_text=detected_text,
                                   file_stream=file_stream)
        else:
            return "No text detected."

    else:
        image_file = request.files['image']
        content = image_file.read()
        client = vision.ImageAnnotatorClient()
        image = vision.Image(content=content)

        response = client.text_detection(image=image)
        texts = response.text_annotations

        if texts:
            detected_text = texts[0].description

            # Create a temporary DOCX file in memory
            doc = Document()
            doc.add_paragraph(detected_text)

            # Save to memory instead of disk
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            return render_template('result.html',
                                   detected_text=detected_text,
                                   file_stream=file_stream)
        else:
            return "No text detected."

#REGISTER PAGE RENDER ROUTE
@app.route("/register")
def register():
    return render_template("register.html")

#LOGIN PAGE RENDER ROUTE
@app.route("/login")
def login():
    return render_template("login.html")

#GO TO DASHBOARD REDIRECT ROUTE
@app.route("/user")
def user():
    if "username" in session:
            username = session.get("username") or "user"
            found_all_users = org_users.query.filter_by(org_username=username).first()
            print_user_pics = found_all_users.id
            found_all_users_id = org_pictures.query.filter_by(user_id=print_user_pics).all()
            return render_template("user.html", users=found_all_users_id,username=username)
    else:
        return redirect(url_for("register"))

#PERFORM PDF OCR
@app.route("/pdf-upload",methods=["POST"])
def pdf_upload():
    pdf = request.files['pdf-submit']

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
        pdf.save(temp_pdf)
        temp_pdf_path = temp_pdf.name  # Path to the temp file

    pdftest = convert_from_path(temp_pdf_path, dpi=300)

    all_text = ""
    for i, img in enumerate(pdftest):
        text = pytesseract.image_to_string(img)
        all_text += f"--- Page {i + 1} ---\n{text}\n\n"

    # CLEANUP
    if os.path.exists(temp_pdf_path):
        os.remove(temp_pdf_path)

    return render_template("result.html", detected_text=all_text)

#IMAGE SHARING ROUTE
@app.route('/share_image', methods=['GET', 'POST'])
def share_image():
    if 'username' not in session:
        flash("You must be logged in to share images.")
        return redirect(url_for("login"))

    username = session.get("username")
    user = org_users.query.filter_by(org_username=username).first()

    if request.method == 'POST':
        # Get the selected image and recipient user
        image_id = request.form.get('image_id')
        recipient_id = request.form.get('recipient_id')

        image = org_pictures.query.filter_by(id=image_id, user_id=user.id).first()
        recipient = org_users.query.filter_by(id=recipient_id).first()

        if image and recipient:
            # Create a new PictureShare record
            new_share = PictureShare(picture_id=image.id, shared_with_user_id=recipient.id)
            db.session.add(new_share)
            db.session.commit()
            flash("Image shared successfully.")
        else:
            flash("Error: Could not find image or recipient.")

        return redirect(url_for('share_image'))

    # Fetch images uploaded by the logged-in user
    user_images = org_pictures.query.filter_by(user_id=user.id).all()

    # Fetch all users to share with
    users = org_users.query.filter(org_users.id != user.id).all()

    return render_template('share_image.html', images=user_images, users=users)

#VIEW RECIEVED IMAGES ROUTE
@app.route('/view_shared_files')
def view_shared_files():
    if 'username' not in session:
        flash("You must be logged in to view shared files.")
        return redirect(url_for("login"))

    username = session.get("username")
    user = org_users.query.filter_by(org_username=username).first()

    if user:
        # Fetch images shared with the user
        shared_images = db.session.query(org_pictures).join(PictureShare).filter(
            PictureShare.shared_with_user_id == user.id).all()

        # Pass the images to the template
        return render_template('view_shared_files.html', shared_images=shared_images)
    else:
        flash("User not found.")
        return redirect(url_for("login"))

#DOWNLOAD RECIEVED IMAGES ROUTE
@app.route('/download_shared_file/<int:image_id>')
def download_shared_file(image_id):
    if 'username' not in session:
        flash("You must be logged in to download shared files.")
        return redirect(url_for("login"))

    username = session.get("username")
    user = org_users.query.filter_by(org_username=username).first()

    if user:
        # Check if this image was shared with the user
        shared = PictureShare.query.filter_by(
            picture_id=image_id,
            shared_with_user_id=user.id
        ).first()

        if shared:
            image = org_pictures.query.filter_by(id=image_id).first()
            if image:
                return send_file(
                    os.path.join(app.config['UPLOAD_FOLDER'], image.image_path),
                    as_attachment=True,
                    download_name=image.image_path
                )
            else:
                flash("Image not found.")
                return redirect(url_for('view_shared_files'))
        else:
            flash("You don't have permission to download this file.")
            return redirect(url_for('view_shared_files'))
    else:
        flash("User not found.")
        return redirect(url_for("login"))

#DATABASE AJAX USERNAME AVAILABILITY CHECK
@app.route('/check_username')
def check_username():
    username = request.args.get('username', '').strip()
    if not username:
        return jsonify({'exists': False})

    # Query the org_users table
    user = org_users.query.filter_by(org_username=username).first()
    exists = user is not None

    return jsonify({'exists': exists})

#<------------------------------------------WORKING CODE ABOVE------------------------------------------------------>
#<------------------------------------------TEST CODE BELOW---------------------------------------------------------->

@app.route('/pdfocr')
def pdfocr():
    return render_template('pdfocr.html')

@app.route('/imageocr')
def imageocr():
    return render_template('imageocr.html')

@app.route('/tessocr')
def tessocr():
    return render_template('tessocr.html')



#<-------------------------------------------END OF TEST CODE-------------------------------------------------------->

#MAIN FUNCTION
if __name__ == "__main__":
    with app.app_context():
        db.create_all()
    app.run(host="0.0.0.0", port=5000, debug=True)
